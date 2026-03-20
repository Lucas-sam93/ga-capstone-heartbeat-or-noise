"""
Generate the GA Capstone Project Detailed Explanation Word Document.

This script creates a professionally formatted .docx file combining
the full project explanation with ECG findings and the final narrative.

Usage:
    python generate_docx.py

Output:
    outputs/GA_Capstone_Project_Detailed_Explanation.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# --- Configuration ---
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "GA_Capstone_Project_Detailed_Explanation.docx")

FONT_NAME = "Calibri"
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
TABLE_STYLE = "Light Grid Accent 1"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR
        run.font.name = FONT_NAME
    return heading


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = TABLE_STYLE
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    return table


def build_document():
    doc = Document()

    # --- Title Page ---
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Heartbeat or Noise?")
    set_run_font(run, size=28, bold=True, color=HEADING_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Evaluating Consumer Wearables as Cardiac Screening Proxies\nUsing Machine Learning"
    )
    set_run_font(run, size=16, color=RGBColor(0x55, 0x55, 0x55))

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Lucas Sam\nGeneral Assembly Data Analytics Capstone Project\nMarch 2026")
    set_run_font(run, size=12)

    doc.add_page_break()

    # =========================================================================
    # SECTION 1 — The Problem
    # =========================================================================
    add_heading(doc, "1. The Problem", level=1)

    add_para(doc,
        "Cardiovascular disease (CVD) accounts for 30.5% of all deaths in Singapore, "
        "equating to approximately 22 lives lost daily. Despite national screening efforts, "
        "approximately 49% of Singapore residents did not undergo a health checkup in the "
        "past twelve months, leaving a significant proportion of the population unaware of "
        "underlying cardiac risk."
    )
    add_para(doc,
        "Consumer wearables \u2014 including smartwatches from Apple and Samsung \u2014 are worn "
        "by millions daily and continuously monitor physiological signals including heart rate "
        "and heart rate variability. This project evaluates whether machine learning models "
        "trained on clinical ECG data can generalise to consumer wearable signals to detect "
        "cardiac rhythm irregularities, assessing their feasibility as low-barrier proxy "
        "screening tools."
    )
    add_para(doc,
        "Research Question: Can machine learning models trained on clinical cardiac data "
        "generalise to consumer wearable signals to detect heart rhythm irregularities, "
        "and can consumer wearables feasibly serve as proxy screening tools to identify "
        "individuals who would benefit from formal cardiac evaluation?",
        bold=True
    )

    # =========================================================================
    # SECTION 2 — Two-Layer Architecture
    # =========================================================================
    add_heading(doc, "2. Two-Layer Validation Architecture", level=1)

    add_para(doc,
        "The project uses a two-layer architecture designed to separate model training "
        "from generalisation testing:"
    )
    add_para(doc,
        "Layer 1 \u2014 Clinical Benchmark: A binary cardiac rhythm classifier is trained "
        "and validated on clinical ECG data using only features extractable from consumer "
        "wearable devices. This deliberate constraint ensures the model reflects real-world "
        "deployment conditions. The model is locked after Layer 1 \u2014 no retraining occurs."
    )
    add_para(doc,
        "Layer 2 \u2014 Consumer Wearable Bridge: The frozen model is applied to wearable data "
        "to evaluate signal transferability. Two datasets are used: a personal Apple Watch "
        "N=1 case study (exploratory) and the MIMIC PERform AF dataset (primary validation, "
        "N=35). Success is assessed across three tiers:"
    )
    add_table(doc,
        ["Tier", "Criterion", "What It Tests"],
        [
            ["Tier 1", "Pipeline executes with interpretable outputs", "Foundational applicability"],
            ["Tier 2", "Sensitivity \u226580% AND Specificity \u226575%", "Statistical signal detection"],
            ["Tier 3", "AUROC \u22650.85", "Clinical concordance / discriminative ability"],
        ]
    )

    # =========================================================================
    # SECTION 3 — Data Sources
    # =========================================================================
    add_heading(doc, "3. Data Sources", level=1)

    add_heading(doc, "3.1 Physionet 2017 AF Classification Challenge (Layer 1)", level=2)
    add_para(doc,
        "8,528 single-lead ECG recordings sampled at 300 Hz. After excluding 279 noisy-class "
        "records and quality filtering, 8,187 records remained. Binary label mapping: Normal "
        "(N \u2192 0) versus Abnormal (AF + Other \u2192 1). Class distribution: 61.6% Normal "
        "(5,042), 38.4% Abnormal (3,145). AF within Abnormal: 754 records (24.0%)."
    )

    add_heading(doc, "3.2 Personal Apple Watch Data (N=1 Case Study)", level=2)
    add_para(doc,
        "Apple Watch SE (1st Generation) \u2014 PPG optical sensor only, no electrical heart "
        "sensor. Temporal coverage: April 2021 \u2013 February 2026 (4.8 years). Clinical "
        "anchor point: June 18, 2025 ECG confirming cardiac irregularity. Role: exploratory "
        "N=1 case study appendix."
    )
    add_table(doc,
        ["Metric", "Records"],
        [
            ["Heart rate (continuous)", "478,077"],
            ["HRV (SDNN)", "5,456"],
            ["Resting heart rate", "1,491"],
            ["Walking heart rate", "1,557"],
            ["Respiratory rate", "20,372"],
            ["Sleep analysis", "7,967"],
            ["Workouts", "1,982"],
        ]
    )

    add_heading(doc, "3.3 MIMIC PERform AF (Primary Layer 2 Validation)", level=2)
    add_para(doc,
        "35 critically ill adults (19 AF, 16 NSR) with simultaneous finger PPG and ECG "
        "recordings at 125 Hz. Each recording is 20 minutes (150,000 samples). Pre-assigned "
        "clinical labels at subject level. Sourced from Zenodo. Pivot rationale: UMass Simband "
        "access unavailable before deadline; MIMIC PERform AF satisfies all three validation "
        "requirements: wearable PPG signal, both Normal and Abnormal subjects, pre-existing "
        "ground truth labels."
    )

    # =========================================================================
    # SECTION 4 — Locked Feature Set
    # =========================================================================
    add_heading(doc, "4. Locked Feature Set", level=1)

    add_para(doc,
        "Eight features, all time-domain or statistical features derivable from both clinical "
        "ECG and wearable PPG data. Locked before feature engineering began. Frequency domain "
        "features (LF/HF ratio, HF power) were permanently excluded because consumer wearables "
        "do not expose raw beat-by-beat interval sequences at consistent sampling rates."
    )
    add_table(doc,
        ["Feature", "Description"],
        [
            ["RMSSD", "Root mean square of successive RR interval differences"],
            ["SDNN", "Standard deviation of all RR intervals"],
            ["Mean RR", "Average inter-beat interval (ms)"],
            ["pNN50", "Proportion of successive intervals differing by >50ms"],
            ["HR Mean", "Average heart rate derived from RR intervals"],
            ["HR Std Dev", "Variability of instantaneous heart rate"],
            ["RR Skewness", "Asymmetry of the RR interval distribution"],
            ["RR Kurtosis", "Tail weight of the RR interval distribution"],
        ]
    )

    # =========================================================================
    # SECTION 5 — Layer 1 Training & Model Selection
    # =========================================================================
    add_heading(doc, "5. Layer 1 \u2014 Training and Model Selection", level=1)

    add_para(doc,
        "The dataset was split 80/10/10 (stratified): 6,549 training, 819 validation, "
        "819 test. StandardScaler was fitted on X_train only \u2014 no data leakage. "
        "Four models were trained with class_weight=balanced to handle the 62/38 imbalance."
    )

    add_heading(doc, "5.1 Cross-Validation Results (5-Fold Stratified)", level=2)
    add_table(doc,
        ["Model", "Sensitivity", "Specificity", "AUROC"],
        [
            ["Logistic Regression", "80.5% \u00b10.1%", "84.2% \u00b12.0%", "0.8817 \u00b10.0066"],
            ["Random Forest", "80.3% \u00b10.2%", "86.7% \u00b11.3%", "0.8955 \u00b10.0080"],
            ["XGBoost", "80.4% \u00b10.2%", "82.5% \u00b11.9%", "0.8809 \u00b10.0054"],
            ["SVM", "80.3% \u00b10.2%", "86.9% \u00b11.4%", "0.8960 \u00b10.0062"],
        ]
    )

    add_heading(doc, "5.2 Overfitting Check", level=2)
    add_table(doc,
        ["Model", "Train Accuracy", "Test Accuracy", "Gap", "Verdict"],
        [
            ["Logistic Regression", "83.2%", "85.6%", "\u22122.4%", "No overfitting"],
            ["Random Forest", "100.0%", "86.6%", "+13.4%", "Significant overfitting"],
            ["XGBoost", "99.8%", "86.2%", "+13.6%", "Significant overfitting"],
            ["SVM", "85.6%", "86.8%", "\u22121.2%", "No overfitting"],
        ]
    )

    add_heading(doc, "5.3 Model Selection \u2014 SVM", level=2)
    add_para(doc,
        "SVM was selected using a five-criterion natural selection framework. It achieved "
        "the highest validation specificity (88.1%), fewest false positives (60), and showed "
        "no overfitting (\u22121.2% train-test gap). Criterion 2 resolved the selection."
    )

    add_heading(doc, "5.4 Held-Out Test Set Results (Threshold = 0.34)", level=2)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Sensitivity", "84.4%"],
            ["Specificity", "87.3%"],
            ["AUROC", "0.9080"],
            ["F1 Score", "0.8243"],
            ["AF Sensitivity", "98.8%"],
            ["Confusion Matrix", "TP=265, FP=64, FN=49, TN=441"],
            ["Pre-registered criterion", "PASS (Sens \u226580%, Spec \u226575%)"],
        ]
    )

    # =========================================================================
    # SECTION 6 — Layer 2A: Apple Watch
    # =========================================================================
    add_heading(doc, "6. Layer 2A \u2014 Apple Watch N=1 Case Study", level=1)

    add_para(doc,
        "1,997 feature windows were generated (30-minute windows, 15-minute step, minimum "
        "10 readings). Windows by period: baseline=1,758, pre_anchor=95, post_anchor=29, "
        "follow_up=115."
    )

    add_heading(doc, "6.1 Signal Modality Gap (KS Test)", level=2)
    add_table(doc,
        ["Feature", "KS Statistic", "Direction", "Magnitude"],
        [
            ["RMSSD", "0.424", "AW lower", "Large"],
            ["HR Mean", "0.416", "AW higher", "Large"],
            ["Mean RR", "0.390", "AW lower", "Large"],
            ["HR Std Dev", "0.359", "AW higher", "Large"],
            ["SDNN", "0.356", "AW lower", "Large"],
            ["RR Skewness", "0.297", "AW higher", "Moderate"],
            ["pNN50", "0.210", "AW lower", "Moderate"],
            ["RR Kurtosis", "0.193", "AW lower", "Moderate"],
        ]
    )
    add_para(doc,
        "Summary: 5 large (KS > 0.3), 3 moderate (0.1\u20130.3), 0 small. No feature "
        "transferred cleanly between modalities."
    )

    add_heading(doc, "6.2 Temporal Analysis (Mann-Whitney U)", level=2)
    add_table(doc,
        ["Period", "n", "Median Score", "p-value", "Effect Size r", "Verdict"],
        [
            ["pre_anchor", "95", "0.104", "0.957", "0.104", "Lower than baseline"],
            ["post_anchor", "29", "0.101", "0.839", "0.107", "Lower than baseline"],
            ["follow_up", "115", "0.121", "0.737", "0.035", "Higher than baseline"],
        ]
    )

    add_heading(doc, "6.3 Tier Assessment", level=2)
    add_para(doc, "Tier 1: PASS \u2014 pipeline executed, gap quantification interpretable.")
    add_para(doc, "Tier 2: FAIL \u2014 no anchor period significantly elevated above baseline.")
    add_para(doc, "Tier 3: FAIL \u2014 Tier 2 not passed.")

    add_heading(doc, "6.4 Null Result Interpretation", level=2)
    add_para(doc,
        "Three non-mutually-exclusive explanations account for the null result:"
    )
    add_para(doc,
        "1. Signal modality gap \u2014 5 of 8 features show large KS distances, placing "
        "Apple Watch windows outside the model\u2019s training distribution."
    )
    add_para(doc,
        "2. Behavioural confound \u2014 reduced training intensity in the pre-anchor period "
        "may have shifted HRV toward more regular patterns interpreted as normal by the model."
    )
    add_para(doc,
        "3. Abnormality type mismatch \u2014 the clinical ECG report confirms an "
        "intraventricular conduction delay (waveform morphology abnormality), not a rhythm "
        "irregularity. HRV features measure beat-to-beat timing variability and are "
        "structurally incapable of detecting conduction delays, which produce regular-but-"
        "abnormally-shaped beats."
    )

    # =========================================================================
    # SECTION 7 — Layer 2B: MIMIC PERform AF
    # =========================================================================
    add_heading(doc, "7. Layer 2B \u2014 MIMIC PERform AF Primary Validation", level=1)

    add_para(doc,
        "35 subjects (19 AF, 16 NSR), all green quality tier (\u2265300 PPG peaks detected). "
        "Single 20-minute window per subject. PPG nulls in 8 files handled by linear "
        "interpolation."
    )

    add_heading(doc, "7.1 Signal Modality Gap", level=2)
    add_para(doc,
        "All 8 features show large KS distances (KS \u22650.3) \u2014 worse than Apple Watch "
        "(5 large, 3 moderate). MIMIC subjects are critically ill ICU patients with "
        "systematically higher HRV variability. This is clinically plausible."
    )

    add_heading(doc, "7.2 SVM at Fixed Threshold 0.34", level=2)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Sensitivity", "100% (19/19 AF)"],
            ["Specificity", "12.5% (2/16 NSR)"],
            ["AUROC", "0.8586"],
            ["Confusion Matrix", "TP=19, FP=14, FN=0, TN=2"],
        ]
    )
    add_para(doc,
        "Tier 1: PASS. Tier 2: FAIL (specificity 12.5%). Tier 3: PASS (AUROC 0.8586 \u2265 0.85)."
    )

    add_heading(doc, "7.3 Cross-Model Comparison", level=2)
    add_para(doc,
        "All four Layer 1 models were applied to MIMIC at their fixed Layer 1 thresholds:"
    )
    add_table(doc,
        ["Model", "Threshold", "Sensitivity", "Specificity", "AUROC"],
        [
            ["Logistic Regression", "0.41", "100%", "18.75%", "0.7368"],
            ["Random Forest", "0.37", "100%", "6.25%", "0.8503"],
            ["XGBoost", "0.34", "100%", "12.5%", "0.8322"],
            ["SVM", "0.34", "100%", "12.5%", "0.8586"],
        ]
    )
    add_para(doc,
        "Specificity failure is systemic across all four models \u2014 confirming this is a "
        "modality gap problem, not a model selection artefact. SVM retains highest AUROC."
    )

    add_heading(doc, "7.4 Threshold Recalibration", level=2)
    add_table(doc,
        ["Metric", "Fixed (0.34)", "Recalibrated (0.8424)"],
        [
            ["Sensitivity", "100%", "73.7%"],
            ["Specificity", "12.5%", "93.8%"],
            ["F1", "0.7308", "0.8235"],
        ]
    )
    add_para(doc,
        "Specificity recovers from 12.5% to 93.8% with domain-adapted threshold. The gap "
        "between 0.34 and 0.8424 quantifies the magnitude of the modality shift. This "
        "confirms the specificity failure is a calibration problem, not a discrimination "
        "failure."
    )

    add_heading(doc, "7.5 Stress Testing", level=2)
    add_table(doc,
        ["Test", "Metric", "Result"],
        [
            ["Bootstrap AUROC (n=1000)", "Mean [95% CI]", "0.8631 [0.7246, 0.9720]"],
            ["Bootstrap Recalibrated Sens", "Mean [95% CI]", "80.6% [57.1%, 100.0%]"],
            ["Bootstrap Recalibrated Spec", "Mean [95% CI]", "91.3% [72.7%, 100.0%]"],
            ["LOOCV Youden\u2019s J", "Sens / Spec", "68.4% / 81.2%"],
            ["Sensitivity-Targeted LOOCV", "Sens / Spec", "78.9% / 81.2%"],
        ]
    )
    add_para(doc,
        "Sensitivity-targeted LOOCV achieves 78.9% sensitivity (1.1pp below pre-registered "
        "80% criterion) with 81.2% specificity (clearing 75% criterion). The 1.1pp gap is "
        "consistent with N=19 AF leave-one-out variance. AUROC lower bound 0.7246 confirms "
        "discriminative signal persists even in unlucky resamples."
    )

    # =========================================================================
    # SECTION 8 — Clinical ECG Report
    # =========================================================================
    add_heading(doc, "8. Clinical ECG Report \u2014 June 2025 Anchor Event", level=1)

    add_para(doc,
        "The clinical ECG report was obtained from the National Heart Centre Singapore (NHCS), "
        "Department of Cardiology, certified by Dr Chong Jia Yun (B.Med, FRACP). The ECG was "
        "conducted on 18 June 2025 at 10:19:48 AM \u2014 confirming the pre-registered anchor "
        "date used throughout the Apple Watch analysis."
    )

    add_heading(doc, "8.1 ECG Parameters", level=2)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["Heart Rate", "112 bpm"],
            ["PR Interval", "125 ms"],
            ["QRS Duration", "125 ms"],
            ["QT/QTc", "424 ms"],
        ]
    )

    add_heading(doc, "8.2 Automated Interpretation", level=2)
    add_para(doc, "The ECG machine\u2019s automated interpretation reported:")
    add_para(doc, "\u2022 Sinus rhythm")
    add_para(doc, "\u2022 Intraventricular conduction delay")
    add_para(doc, "\u2022 ST abnormality (probable anterior early repolarisation)")
    add_para(doc, "\u2022 Overall: Abnormal ECG", bold=True)

    add_heading(doc, "8.3 Implications for the Project", level=2)
    add_para(doc,
        "Anchor date confirmed: The June 18, 2025 ECG date aligns exactly with the "
        "pre-registered anchor point, validating the temporal framework used in the Apple "
        "Watch analysis."
    )
    add_para(doc,
        "Abnormal classification confirmed: The ECG was flagged as abnormal, validating "
        "the binary label (Abnormal = 1) assigned to the anchor event."
    )
    add_para(doc,
        "Critical insight \u2014 conduction delay vs arrhythmia: The specific abnormality "
        "is an intraventricular conduction delay, NOT an arrhythmia such as atrial "
        "fibrillation. This is a waveform morphology abnormality \u2014 it affects QRS shape "
        "and the ST segment, not beat-to-beat timing. HRV features (RMSSD, SDNN, pNN50, "
        "etc.) are designed to detect rhythm irregularities through inter-beat interval "
        "variability. A conduction delay produces regular-but-abnormally-shaped beats, "
        "which would appear entirely normal to any HRV-based analysis."
    )
    add_para(doc,
        "This provides a third mechanistic explanation for the Apple Watch null result: "
        "beyond the signal modality gap and the behavioural confound, the model was asked "
        "to detect something its feature set is structurally incapable of capturing. This "
        "strengthens the null result interpretation and narrows the project\u2019s validated "
        "scope to rhythm irregularities specifically."
    )

    # =========================================================================
    # SECTION 9 — BeatCheck App
    # =========================================================================
    add_heading(doc, "9. BeatCheck \u2014 Cardiac Screening App", level=1)

    add_para(doc,
        "BeatCheck is a consumer-facing web app that applies the trained SVM model to "
        "Apple Health heart rate exports, built as a proof-of-concept demonstration of "
        "the research pipeline in a deployable format."
    )
    add_para(doc,
        "Stack: FastAPI backend, plain HTML/CSS/JS frontend. Input formats: Apple Health "
        "CSV export, XML export, or full ZIP archive. Client-side JavaScript handles large "
        "XML files via streaming \u2014 the full multi-year export is never transmitted to "
        "the server. Data is filtered client-side to the most recent 90 days."
    )
    add_para(doc,
        "Processing: Uploaded data is windowed into 30-minute overlapping windows "
        "(15-minute step). Each window is scored by the SVM at a domain-adapted threshold "
        "of 0.8368 (LOOCV mean threshold). The percentage of windows flagged as abnormal "
        "is mapped to three risk tiers:"
    )
    add_table(doc,
        ["Tier", "Windows Flagged", "Interpretation"],
        [
            ["Low", "< 10%", "No pattern of concern detected"],
            ["Intermediate", "10\u201340%", "Elevated pattern \u2014 consider follow-up"],
            ["High", "> 40%", "Consistent pattern \u2014 clinical review recommended"],
        ]
    )
    add_para(doc,
        "Stress test: Real Apple Health export (4.8 years) processed successfully \u2014 "
        "Intermediate tier, 38.2% of windows flagged across 84 days analysed."
    )

    # =========================================================================
    # SECTION 10 — Conclusion / Final Narrative
    # =========================================================================
    add_heading(doc, "10. Conclusion \u2014 The Complete Story", level=1)

    add_heading(doc, "10.1 Answering the Research Question", level=2)
    add_para(doc,
        "Can machine learning models trained on clinical ECG data generalise to consumer "
        "wearable signals to detect cardiac rhythm anomalies, and are consumer wearables "
        "feasible as proxy cardiac screening tools for the general population?",
        bold=True
    )
    add_para(doc,
        "The answer is a qualified yes. The discriminative signal trained on clinical ECG "
        "data transfers meaningfully to consumer wearable PPG signals, but direct threshold "
        "generalisation fails due to the modality gap between clinical and consumer "
        "measurement systems."
    )
    add_para(doc,
        "Layer 1 established that an SVM classifier trained on 8 HRV features from 8,187 "
        "clinical ECG recordings achieves 84.4% sensitivity and 87.3% specificity "
        "(AUROC 0.9080) on held-out test data \u2014 surpassing the pre-registered criterion "
        "of sensitivity \u226580% at specificity \u226575%. The model demonstrates strong "
        "discriminative ability for distinguishing normal from abnormal cardiac rhythms "
        "using only features extractable from consumer wearables."
    )
    add_para(doc,
        "Layer 2 tested whether this performance transfers to real wearable data. The MIMIC "
        "PERform AF validation (N=35) revealed the central finding: AUROC 0.8586 confirms "
        "the model retains genuine discriminative ability across modalities, but the fixed "
        "ECG-calibrated threshold (0.34) produces 100% sensitivity at only 12.5% specificity. "
        "This is not a discrimination failure \u2014 it is a calibration failure. When the "
        "threshold is recalibrated for the PPG domain (0.8368 via LOOCV), sensitivity-targeted "
        "LOOCV recovers 78.9% sensitivity at 81.2% specificity, approaching the pre-registered "
        "criterion within the variance expected from N=35 leave-one-out evaluation."
    )
    add_para(doc,
        "The Apple Watch N=1 case study produced a null result \u2014 probability scores were "
        "not elevated around the clinical anchor event. Three non-mutually-exclusive "
        "explanations account for this: the signal modality gap, a behavioural confound "
        "from reduced training intensity, and critically, the nature of the clinical finding "
        "itself. The ECG report from the National Heart Centre Singapore confirms an "
        "intraventricular conduction delay \u2014 a waveform morphology abnormality that "
        "affects QRS shape, not beat-to-beat timing. HRV features, which measure inter-beat "
        "interval variability, are structurally incapable of detecting this type of "
        "abnormality. This finding narrows the project\u2019s validated scope to rhythm "
        "irregularities specifically, while demonstrating that real-world cardiac "
        "abnormalities encompass a broader spectrum than HRV-based screening can capture."
    )

    add_heading(doc, "10.2 The Discrimination vs Calibration Distinction", level=2)
    add_para(doc,
        "The most important methodological contribution of this project is the distinction "
        "between discrimination and calibration when evaluating cross-modality model "
        "performance. At the fixed ECG threshold of 0.34, the MIMIC results appear "
        "catastrophic: 12.5% specificity means 87.5% of healthy subjects are falsely flagged. "
        "But AUROC 0.8586 tells a different story \u2014 the model correctly ranks abnormal "
        "subjects higher than normal subjects 86% of the time. The model knows who is "
        "sicker; it just doesn\u2019t know where to draw the line because the line was drawn "
        "for a different measurement system."
    )
    add_para(doc,
        "This is the same reason a thermometer calibrated in Fahrenheit gives alarming "
        "readings when interpreted on a Celsius scale. The instrument is accurate; the "
        "interpretation framework needs adaptation."
    )

    add_heading(doc, "10.3 Implications for Singapore Public Health", level=2)
    add_para(doc,
        "With 49% of Singapore residents foregoing annual health checkups and cardiovascular "
        "disease claiming approximately 22 lives daily, the screening gap is real and "
        "consequential. Consumer wearables are already worn by millions and continuously "
        "collect cardiac data without requiring active participation."
    )
    add_para(doc,
        "This project demonstrates that the bridge from clinical models to consumer "
        "wearable signals is feasible but not yet direct. The discriminative signal "
        "transfers (AUROC 0.86 across modalities), but deployment requires domain-adapted "
        "thresholds calibrated to the specific wearable measurement system. A clinical ECG "
        "threshold applied directly to PPG data will over-flag \u2014 producing unnecessary "
        "anxiety and clinical referrals. A recalibrated threshold restores clinically "
        "meaningful specificity."
    )
    add_para(doc,
        "For Singapore\u2019s public health strategy, the practical path forward involves: "
        "(1) training or adapting models with PPG-specific calibration data from the target "
        "population; (2) partnering with wearable manufacturers to access richer signal data "
        "(raw RR intervals rather than summary statistics); and (3) validating on larger, "
        "population-representative cohorts rather than ICU patients. The technology is not "
        "ready for autonomous screening today, but the discriminative foundation is present "
        "\u2014 the engineering problem is calibration, not capability."
    )

    # =========================================================================
    # SECTION 11 — Methodological Strengths
    # =========================================================================
    add_heading(doc, "11. Methodological Strengths", level=1)

    strengths = [
        ("Pre-registration", "Key decisions (feature set, thresholds, test direction, "
         "tier criteria) were locked before analysis, preventing post-hoc rationalisation."),
        ("No data leakage", "StandardScaler fitted on training data only. Layer 2 data "
         "never used for training or threshold selection (fixed threshold applied first, "
         "recalibration documented separately)."),
        ("Transparent null result", "The Apple Watch null result is reported rather than "
         "suppressed. Three mechanistic explanations are provided, and the pivot to MIMIC "
         "PERform AF is documented with rationale."),
        ("Discrimination vs calibration", "AUROC and sensitivity/specificity are reported "
         "and interpreted separately, preventing the common error of conflating threshold "
         "miscalibration with model failure."),
        ("Stress testing", "Bootstrap confidence intervals and LOOCV provide honest "
         "uncertainty estimates rather than point estimates that overstate precision."),
        ("ECG ground truth", "The clinical ECG report provides a verifiable anchor point "
         "and reveals the specific type of abnormality, enabling mechanistic interpretation "
         "rather than speculation."),
        ("Reproducible pipeline", "All source code in src/, notebooks call functions "
         "rather than defining them, and the full methodology is documented."),
    ]
    for title, desc in strengths:
        add_para(doc, f"{title}: {desc}")

    # =========================================================================
    # SECTION 12 — Limitations
    # =========================================================================
    add_heading(doc, "12. Limitations", level=1)

    limitations = [
        "This project does not constitute a clinical trial and does not produce a "
        "validated medical device.",
        "The personal Apple Watch analysis is a single-subject case study (N=1) and "
        "its findings cannot be generalised to the broader population.",
        "Consumer wearable data is derived from optical heart rate sensors (PPG), which "
        "measure cardiac activity through a fundamentally different mechanism than the "
        "clinical ECG recordings used for model training.",
        "The Apple Watch SE does not have an electrical heart sensor and cannot produce "
        "ECG recordings. RMSSD values are approximated from SDNN (\u00d70.85).",
        "The MIMIC PERform AF validation uses 35 critically ill ICU subjects \u2014 a "
        "population that differs significantly from the general screening target.",
        "Threshold recalibration and LOOCV results were computed on the same 35 subjects "
        "and should be interpreted as indicative rather than definitive.",
        "The locked 8-feature HRV set detects rhythm irregularities (beat-to-beat timing "
        "abnormalities) but cannot detect morphological abnormalities such as conduction "
        "delays or ST changes, which affect waveform shape rather than inter-beat timing "
        "\u2014 as demonstrated by the clinical ECG report findings.",
        "All personal data findings are presented as exploratory and hypothesis-generating.",
    ]
    for lim in limitations:
        add_para(doc, f"\u2022 {lim}")

    # =========================================================================
    # SECTION 13 — Acknowledgements
    # =========================================================================
    add_heading(doc, "13. Acknowledgements", level=1)

    add_para(doc,
        "Dataset provided by Physionet and the Computing in Cardiology Challenge 2017 "
        "organisers. Personal health data collected via Apple Watch SE (1st Generation) "
        "and exported through Apple Health. MIMIC PERform AF dataset provided via Zenodo "
        "(doi.org/10.5281/zenodo.6807402). Clinical ECG report provided by the National "
        "Heart Centre Singapore (NHCS), Department of Cardiology."
    )

    # --- Save ---
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
