"""
Generate the GA Capstone Project Technical Report in DOCX format.

This script creates a professionally formatted technical report structured
around the GA rubric: Problem / Data / Patterns / Model / Recommendations.
Figures from outputs/figures/ are embedded inline.

Usage:
    python scripts/generate_report_docx.py

Output:
    docs/technical_report.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(PROJECT_ROOT, "outputs", "figures")
DOCS_DIR = os.path.join(PROJECT_ROOT, "docs")
OUTPUT_FILE = os.path.join(DOCS_DIR, "technical_report.docx")

FONT_NAME = "Calibri"
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)   # Dark navy
SUBHEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)
TABLE_STYLE = "Light Grid Accent 1"

os.makedirs(DOCS_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_run_font(run, size=11, bold=False, color=None, italic=False):
    """Apply consistent font styling to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """Add a styled heading with the project colour scheme."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR
        run.font.name = FONT_NAME
    return heading


def add_para(doc, text, bold=False, size=11, italic=False, indent=False):
    """Add a paragraph with consistent font styling."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_table(doc, headers, rows):
    """Add a formatted table with header row and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = TABLE_STYLE
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    return table


def add_figure(doc, image_path, caption, width_inches=5.5):
    """Embed a figure with a caption. Skips gracefully if file not found."""
    if not os.path.exists(image_path):
        add_para(
            doc,
            f"[Figure not found: {os.path.basename(image_path)}]",
            italic=True,
        )
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def add_page_footer(doc):
    """Add a page number footer to all sections."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run("Page ")
    set_run_font(run, size=9)
    # Insert page number field
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)
    run2 = footer_para.add_run(" | Heart on Your Wrist \u2014 Lucas Sam \u2014 March 2026")
    set_run_font(run2, size=9)


def fig(name):
    """Return the full path for a figure in outputs/figures/."""
    return os.path.join(FIGURES_DIR, name)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()
    add_page_footer(doc)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Heart on Your Wrist")
    set_run_font(run, size=28, bold=True, color=HEADING_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Evaluating Consumer Wearables as Cardiac Screening Proxies\n"
        "Using Machine Learning"
    )
    set_run_font(run, size=16, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run(
        "Lucas Sam\n"
        "General Assembly Data Analytics Capstone Project\n"
        "March 2026"
    )
    set_run_font(run, size=12)

    doc.add_paragraph()

    rq = doc.add_paragraph()
    rq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rq.add_run(
        "Research Question: Can machine learning models trained on clinical ECG data\n"
        "generalise to consumer wearable signals to detect cardiac rhythm anomalies,\n"
        "and are consumer wearables feasible as proxy cardiac screening tools\n"
        "for the general population?"
    )
    set_run_font(run, size=11, italic=True, color=RGBColor(0x44, 0x44, 0x44))

    doc.add_page_break()

    # =========================================================================
    # SECTION 1 \u2014 Problem, Goals & Audience
    # =========================================================================
    add_heading(doc, "1. Problem, Goals & Audience", level=1)

    add_para(doc,
        "Cardiovascular disease (CVD) accounts for 30.5% of all deaths in Singapore, "
        "equating to approximately 22 lives lost daily. Despite national screening programmes, "
        "approximately 49% of Singapore residents did not undergo a health checkup in the past "
        "twelve months, leaving a significant proportion of the population unaware of underlying "
        "cardiac risk."
    )
    add_para(doc,
        "Consumer wearables \u2014 including smartwatches from Apple and Samsung \u2014 are worn "
        "by millions daily and continuously monitor physiological signals including heart rate "
        "and heart rate variability (HRV). Unlike a hospital visit, wearable monitoring is "
        "passive, always-on, and requires no scheduling or cost. If consumer wearable signals "
        "can reliably flag cardiac rhythm irregularities with sufficient sensitivity, they "
        "represent a scalable, low-barrier first-line screening intervention for an "
        "under-screened population."
    )

    add_heading(doc, "1.1 Project Goals", level=2)
    add_para(doc,
        "Goal 1 \u2014 Clinical Benchmark: Train and evaluate a binary cardiac rhythm classifier "
        "on clinical ECG data, using only features extractable from consumer wearables. "
        "Success criterion: sensitivity \u226580%, specificity \u226575% on held-out test data.",
        indent=True
    )
    add_para(doc,
        "Goal 2 \u2014 Generalisation Test: Apply the frozen model to consumer wearable signals "
        "(Apple Watch PPG and MIMIC PERform AF finger PPG) and measure how much discriminative "
        "performance is retained across the modality gap.",
        indent=True
    )
    add_para(doc,
        "Goal 3 \u2014 Proof of Concept: Build a deployable app (BeatCheck) that accepts "
        "Apple Health exports and returns a risk tier using the trained model.",
        indent=True
    )

    add_heading(doc, "1.2 Target Audience", level=2)
    add_para(doc,
        "Primary audience: Singapore residents who own consumer wearables and are among "
        "the 49% who do not attend regular health checkups. Secondary audience: public health "
        "policymakers evaluating technology-assisted screening interventions. "
        "BeatCheck is explicitly designed as a screening flag \u2014 not a diagnostic tool \u2014 "
        "that directs users toward formal clinical evaluation."
    )

    add_heading(doc, "1.3 Two-Layer Validation Framework", level=2)
    add_para(doc,
        "The project uses a pre-registered two-layer architecture to separate model "
        "training from generalisation testing. Layer 1 trains the model on clinical ECG data. "
        "Layer 2 tests the frozen model on consumer wearable data. No retraining or "
        "threshold adjustment occurs between layers at the point of primary evaluation. "
        "Three success tiers were pre-registered:"
    )
    add_table(doc,
        ["Tier", "Criterion", "What It Tests"],
        [
            ["Tier 1", "Pipeline executes with interpretable outputs", "Foundational applicability"],
            ["Tier 2", "Sensitivity \u226580% AND Specificity \u226575%", "Statistical signal detection"],
            ["Tier 3", "AUROC \u22650.85", "Clinical concordance / discriminative ability"],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 2 \u2014 Data Sources & Data Dictionary
    # =========================================================================
    add_heading(doc, "2. Data Sources & Data Dictionary", level=1)

    add_heading(doc, "2.1 Physionet 2017 AF Classification Challenge (Layer 1 Training)", level=2)
    add_para(doc,
        "8,528 single-lead ECG recordings sampled at 300 Hz, from the PhysioNet Computing "
        "in Cardiology Challenge 2017. After excluding 279 noisy/unclassifiable records and "
        "applying a feature quality filter, 8,187 records were retained (0.75% failure rate). "
        "Binary label mapping: Normal (N \u2192 0) versus Abnormal (AF + Other \u2192 1). "
        "Class distribution: 61.6% Normal (5,042 records), 38.4% Abnormal (3,145 records). "
        "AF within the Abnormal class: 754 records (24.0%)."
    )

    add_figure(
        doc,
        fig("class_distribution_2026.png"),
        "Figure 1. Class distribution in the Physionet 2017 training dataset after "
        "noisy-class exclusion and quality filtering."
    )

    add_figure(
        doc,
        fig("pipeline_flowchart_2026.png"),
        "Figure 2. End-to-end project pipeline from raw data through training, "
        "validation, and wearable generalisation."
    )

    add_heading(doc, "2.2 Personal Apple Watch Data (N=1 Case Study)", level=2)
    add_para(doc,
        "Apple Watch SE, 1st Generation \u2014 PPG optical sensor only. No electrical heart "
        "sensor. No ECG recordings. Temporal coverage: April 2021 \u2013 February 2026 (4.8 years). "
        "Clinical anchor point: June 18, 2025 \u2014 confirmed by ECG report from the National "
        "Heart Centre Singapore (NHCS) showing an Abnormal ECG (intraventricular conduction "
        "delay, ST abnormality). Role: exploratory N=1 case study appendix."
    )
    add_table(doc,
        ["Metric", "Records", "Notes"],
        [
            ["Heart rate (continuous)", "478,077", "1 artifact removed"],
            ["HRV (SDNN)", "5,456", "29 outliers removed"],
            ["Resting heart rate", "1,491", "Daily summaries"],
            ["Walking heart rate", "1,557", "Daily summaries"],
            ["Respiratory rate", "20,372", "Sleep-derived"],
            ["Sleep analysis", "7,967", "Session-level"],
            ["Workouts", "1,982", "Session-level"],
        ]
    )

    add_heading(doc, "2.3 MIMIC PERform AF \u2014 Primary Layer 2 Validation", level=2)
    add_para(doc,
        "35 critically ill adults (19 AF, 16 NSR) with simultaneous finger PPG and ECG "
        "recordings at 125 Hz for 20 minutes per subject (150,000 samples). Pre-assigned "
        "clinical labels at subject level: AF files = Abnormal (1), Non-AF files = Normal (0). "
        "Sourced from Zenodo (doi.org/10.5281/zenodo.6807402). PPG nulls in 8 of 35 files "
        "handled by linear interpolation (worst case: 1.15% of samples). Pivot from UMass "
        "Simband (access unavailable before deadline) \u2014 MIMIC PERform AF satisfies all "
        "three validation requirements: wearable PPG signal, binary ground truth labels, "
        "no access barriers."
    )

    add_heading(doc, "2.4 Feature Dictionary", level=2)
    add_para(doc,
        "Eight features, all time-domain or statistical, extractable from both clinical ECG "
        "and consumer wearable PPG. Locked before feature engineering. Frequency domain "
        "features (LF/HF ratio, HF power) permanently excluded \u2014 consumer wearables do "
        "not expose raw beat-by-beat interval sequences at consistent sampling rates."
    )
    add_table(doc,
        ["Feature", "Description", "Physionet Source", "MIMIC / Wearable Source"],
        [
            ["RMSSD", "Root mean square of successive RR differences", "RR intervals from ECG", "RR intervals from PPG peaks"],
            ["SDNN", "Standard deviation of all RR intervals", "RR intervals from ECG", "RR intervals from PPG peaks"],
            ["Mean RR", "Average inter-beat interval (ms)", "RR intervals from ECG", "RR intervals from PPG peaks"],
            ["pNN50", "Proportion of successive intervals differing >50ms", "RR intervals from ECG", "RR intervals from PPG peaks"],
            ["HR Mean", "Average heart rate (bpm)", "Derived from RR intervals", "Derived from RR intervals"],
            ["HR Std Dev", "Variability of instantaneous heart rate", "Derived from RR intervals", "Derived from RR intervals"],
            ["RR Skewness", "Asymmetry of the RR interval distribution", "Derived from RR intervals", "Derived from RR intervals"],
            ["RR Kurtosis", "Tail weight of the RR interval distribution", "Derived from RR intervals", "Derived from RR intervals"],
        ]
    )

    add_figure(
        doc,
        fig("feature_descriptions_2026.png"),
        "Figure 3. Visual descriptions of all eight locked features and their "
        "physiological interpretation."
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 3 \u2014 Patterns, Trends & Insights
    # =========================================================================
    add_heading(doc, "3. Patterns, Trends & Insights", level=1)

    add_heading(doc, "3.1 Feature Distributions by Class", level=2)
    add_para(doc,
        "Examining the distribution of all eight HRV features across Normal and Abnormal "
        "classes reveals consistent, interpretable separation. Abnormal recordings show higher "
        "RMSSD, SDNN, pNN50, and HR Std Dev \u2014 reflecting the increased beat-to-beat "
        "irregularity characteristic of atrial fibrillation and other rhythm anomalies. "
        "Mean RR is lower in Abnormal recordings (corresponding to higher average heart rate). "
        "RR Skewness and Kurtosis capture distributional shape differences associated with "
        "erratic inter-beat timing."
    )

    add_figure(
        doc,
        fig("feature_distributions_by_class_2026.png"),
        "Figure 4. Distribution of all eight HRV features, split by Normal (0) "
        "and Abnormal (1) class. Higher RMSSD, SDNN, HR Std Dev indicate rhythm irregularity."
    )

    add_heading(doc, "3.2 Feature Correlations", level=2)
    add_para(doc,
        "A correlation heatmap of the eight features reveals expected relationships: RMSSD "
        "and SDNN are strongly positively correlated (both measure beat-to-beat variability "
        "magnitude). pNN50 correlates with RMSSD (both sensitive to fast beat-to-beat "
        "changes). Mean RR and HR Mean are near-perfect inverse transforms of each other. "
        "RR Skewness and Kurtosis are relatively independent of the interval-based features, "
        "contributing orthogonal information about distributional shape."
    )

    add_figure(
        doc,
        fig("feature_correlation_heatmap_2026.png"),
        "Figure 5. Pearson correlation heatmap of all eight locked HRV features. "
        "RMSSD\u2013SDNN and Mean RR\u2013HR Mean show the strongest within-group correlations."
    )

    add_heading(doc, "3.3 Feature Importance (Random Forest)", level=2)
    add_para(doc,
        "Although SVM was selected as the final model, Random Forest feature importances "
        "are reported as a supplementary interpretability measure. RMSSD ranks first (19.58%), "
        "reflecting its direct sensitivity to AF-related rapid irregular beats. HR Std Dev "
        "ranks second (15.58%), capturing overall heart rate volatility. Mean RR ranks third "
        "(13.32%), reflecting the elevated average heart rate common in Abnormal rhythms. "
        "All eight features contribute meaningfully \u2014 no single feature dominates exclusively."
    )

    add_figure(
        doc,
        fig("feature_importance_2026.png"),
        "Figure 9. Random Forest feature importance scores for all eight HRV features. "
        "Reported as a supplementary interpretability measure; SVM is the selected model."
    )
    add_table(doc,
        ["Rank", "Feature", "Importance"],
        [
            ["1", "RMSSD", "19.58%"],
            ["2", "HR Std Dev", "15.58%"],
            ["3", "Mean RR", "13.32%"],
            ["4", "HR Mean", "12.39%"],
            ["5", "SDNN", "11.87%"],
            ["6", "RR Skewness", "9.61%"],
            ["7", "pNN50", "9.31%"],
            ["8", "RR Kurtosis", "8.32%"],
        ]
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 4 \u2014 Predictive Model
    # =========================================================================
    add_heading(doc, "4. Predictive Model \u2014 Layer 1 Clinical Benchmark", level=1)

    add_para(doc,
        "The dataset was split 80/10/10 (stratified): 6,549 training, 819 validation, "
        "819 held-out test records. StandardScaler was fitted on X_train only \u2014 no data "
        "leakage. Four binary classifiers were trained with class_weight=balanced to address "
        "the 62/38 class imbalance. The scaler and all trained models are saved to "
        "outputs/models/."
    )

    add_heading(doc, "4.1 Cross-Validation Results (5-Fold Stratified)", level=2)
    add_para(doc,
        "All four models achieve similar cross-validated sensitivity (80.3\u201380.5%) "
        "with the class_weight=balanced constraint. Specificity and AUROC differentiate them:"
    )
    add_table(doc,
        ["Model", "Sensitivity", "Specificity", "AUROC"],
        [
            ["Logistic Regression", "80.5% \u00b10.1%", "84.2% \u00b12.0%", "0.8817 \u00b10.0066"],
            ["Random Forest", "80.3% \u00b10.2%", "86.7% \u00b11.3%", "0.8955 \u00b10.0080"],
            ["XGBoost", "80.4% \u00b10.2%", "82.5% \u00b11.9%", "0.8809 \u00b10.0054"],
            ["SVM", "80.3% \u00b10.2%", "86.9% \u00b11.4%", "0.8960 \u00b10.0062"],
        ]
    )

    add_heading(doc, "4.2 Overfitting Check", level=2)
    add_table(doc,
        ["Model", "Train Accuracy", "Test Accuracy", "Gap", "Verdict"],
        [
            ["Logistic Regression", "83.2%", "85.6%", "\u22122.4%", "No overfitting"],
            ["Random Forest", "100.0%", "86.6%", "+13.4%", "Significant overfitting"],
            ["XGBoost", "99.8%", "86.2%", "+13.6%", "Significant overfitting"],
            ["SVM", "85.6%", "86.8%", "\u22121.2%", "No overfitting"],
        ]
    )
    add_para(doc,
        "Random Forest and XGBoost both memorise the training set (100% and 99.8% training "
        "accuracy). Despite competitive test accuracy, this overfitting indicates the "
        "ensemble models have not learnt generalisable decision boundaries. Logistic "
        "Regression and SVM show no overfitting."
    )

    add_heading(doc, "4.3 Model Comparison", level=2)

    add_figure(
        doc,
        fig("model_comparison_scorecard_2026.png"),
        "Figure 6. Model comparison scorecard across five selection criteria: "
        "AUROC, specificity, false positives, overfitting gap, and AF sensitivity."
    )

    add_heading(doc, "4.4 Model Selection \u2014 SVM", level=2)
    add_para(doc,
        "SVM was selected using a five-criterion natural selection framework. It achieved "
        "the highest validation specificity (88.1%), fewest false positives (60), and no "
        "overfitting (\u22121.2% train-test gap). Criterion 2 (specificity) resolved the selection."
    )
    add_para(doc,
        "In a screening context, false positives \u2014 incorrectly flagging healthy individuals "
        "as abnormal \u2014 cause unnecessary anxiety and unnecessary clinical referrals. SVM "
        "minimising false positives while maintaining sensitivity is the correct trade-off "
        "for this application."
    )

    add_heading(doc, "4.5 Held-Out Test Set Results (Threshold = 0.34)", level=2)
    add_para(doc,
        "The classification threshold 0.34 was selected on the validation set to prioritise "
        "sensitivity (screening context). The frozen threshold was applied once to the "
        "held-out test set."
    )
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Sensitivity", "84.4%"],
            ["Specificity", "87.3%"],
            ["AUROC", "0.9080"],
            ["F1 Score", "0.8243"],
            ["AF-specific Sensitivity", "98.8%"],
            ["Confusion Matrix", "TP=265, FP=64, FN=49, TN=441"],
            ["Pre-registered criterion", "PASS (Sens \u226580%, Spec \u226575%)"],
        ]
    )

    add_figure(
        doc,
        fig("roc_curve_svm_layer1_2026.png"),
        "Figure 7. ROC curve for the SVM classifier on the held-out test set. "
        "AUROC = 0.9080. Operating point at threshold 0.34 marked."
    )

    add_figure(
        doc,
        fig("confusion_matrix_svm_layer1_2026.png"),
        "Figure 8. Confusion matrix for SVM at threshold 0.34 on held-out test data. "
        "TP=265, FP=64, FN=49, TN=441."
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 4a \u2014 Apple Watch Case Study
    # =========================================================================
    add_heading(doc, "4a. Consumer Wearable Bridge \u2014 Apple Watch N=1 Case Study", level=1)

    add_para(doc,
        "The personal Apple Watch dataset provides an exploratory N=1 test of whether "
        "the trained model captures temporal changes in cardiac rhythm patterns around a "
        "confirmed clinical event. The clinical anchor point is the June 18, 2025 ECG at "
        "NHCS confirming an Abnormal ECG."
    )
    add_para(doc,
        "1,997 feature windows were generated (30-minute windows, 15-minute step, minimum "
        "10 heart rate readings per window). RMSSD was approximated from Apple Watch SDNN "
        "(\u00d70.85) because raw beat-by-beat intervals are not exposed by the device API. "
        "Windows lacking an HRV record were dropped (no imputation). Window distribution: "
        "baseline=1,758, pre_anchor=95, post_anchor=29, follow_up=115."
    )

    add_heading(doc, "4a.1 Signal Modality Gap (KS Test)", level=2)
    add_para(doc,
        "The Kolmogorov-Smirnov test quantifies distributional distance between Apple Watch "
        "feature values and the Physionet training distribution. All p-values are effectively "
        "zero, indicating the Apple Watch features occupy a different region of feature space "
        "than the training data."
    )
    add_table(doc,
        ["Feature", "KS Statistic", "Direction", "Magnitude"],
        [
            ["RMSSD", "0.424", "AW lower", "Large"],
            ["HR Mean", "0.416", "AW higher", "Large"],
            ["Mean RR", "0.390", "AW lower", "Large"],
            ["HR Std Dev", "0.359", "AW higher", "Large"],
            ["SDNN", "0.356", "AW lower", "Large"],
            ["RR Skewness", "0.297", "AW higher", "Moderate"],
            ["pNN50", "0.210", "AW lower", "Moderate"],
            ["RR Kurtosis", "0.193", "AW lower", "Moderate"],
        ]
    )
    add_para(doc,
        "Summary: 5 large (KS > 0.3), 3 moderate (0.1\u20130.3), 0 small. "
        "No feature transferred cleanly between modalities."
    )

    add_figure(
        doc,
        os.path.join(FIGURES_DIR, "aw_vs_physionet_distributions.png"),
        "Figure \u2014 Apple Watch vs Physionet feature distributions. "
        "KS distances confirm systematic modality gap across all 8 features."
    )

    add_heading(doc, "4a.2 Temporal Analysis (Mann-Whitney U, One-Sided)", level=2)
    add_para(doc,
        "A one-sided Mann-Whitney U test was pre-registered to test whether probability "
        "scores in each anchor period are higher than baseline. The null hypothesis was "
        "not rejected in any period:"
    )
    add_table(doc,
        ["Period", "n", "Median Score", "p-value", "Effect Size r", "Verdict"],
        [
            ["pre_anchor", "95", "0.104", "0.957", "0.104", "Lower than baseline"],
            ["post_anchor", "29", "0.101", "0.839", "0.107", "Lower than baseline"],
            ["follow_up", "115", "0.121", "0.737", "0.035", "Higher than baseline"],
        ]
    )
    add_para(doc,
        "Tier 1: PASS \u2014 pipeline executed, gap quantification interpretable. "
        "Tier 2: FAIL \u2014 no anchor period elevated above baseline. "
        "Tier 3: FAIL \u2014 Tier 2 not met."
    )

    add_heading(doc, "4a.3 Null Result Interpretation", level=2)
    add_para(doc,
        "Three non-mutually-exclusive explanations account for the null result:"
    )
    add_para(doc,
        "1. Signal modality gap: 5 of 8 features show large KS distances (>0.3), placing "
        "Apple Watch windows outside the model\u2019s training distribution. The model "
        "operates on data that does not resemble what it was trained on.",
        indent=True
    )
    add_para(doc,
        "2. Behavioural confound: Reduced physical training intensity in the pre-anchor "
        "period may have shifted HRV toward more regular patterns that the model interprets "
        "as normal, independent of any cardiac pathology.",
        indent=True
    )
    add_para(doc,
        "3. Abnormality type mismatch: The clinical ECG report confirms an intraventricular "
        "conduction delay \u2014 a waveform morphology abnormality affecting QRS shape and ST "
        "segment, not beat-to-beat timing. HRV features are structurally incapable of "
        "detecting conduction delays, which produce regular-but-abnormally-shaped beats. "
        "The model was asked to detect something its feature set cannot capture.",
        indent=True
    )
    add_para(doc,
        "The null result is reported transparently. It motivates the pivot to MIMIC PERform "
        "AF as primary Layer 2 validation, and demonstrates the scope boundary of "
        "HRV-based screening: rhythm irregularities are detectable; morphological "
        "abnormalities are not."
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 4b \u2014 MIMIC PERform AF Validation
    # =========================================================================
    add_heading(doc, "4b. Consumer Wearable Bridge \u2014 MIMIC PERform AF Primary Validation", level=1)

    add_para(doc,
        "35 subjects (19 AF, 16 NSR). All 35 subjects passed the green quality tier "
        "(\u2265300 PPG peaks detected). Single 20-minute window per subject. The frozen "
        "SVM and frozen scaler (transform only, no refit) were applied at the pre-registered "
        "fixed threshold of 0.34."
    )

    add_heading(doc, "4b.1 Signal Modality Gap (KS Test)", level=2)
    add_para(doc,
        "All 8 features show large KS distances (KS \u22650.3) \u2014 worse than Apple Watch "
        "(5 large). MIMIC subjects are critically ill ICU patients whose underlying physiology "
        "produces systematically elevated HRV variability, pushing all features toward "
        "values the model associates with Abnormal."
    )
    add_table(doc,
        ["Feature", "KS Statistic", "Direction"],
        [
            ["RR Skewness", "0.5186", "MIMIC higher"],
            ["pNN50", "0.4739", "MIMIC higher"],
            ["RR Kurtosis", "0.4121", "MIMIC higher"],
            ["RMSSD", "0.3747", "MIMIC higher"],
            ["Mean RR / HR Mean", "0.3853", "MIMIC mixed"],
            ["HR Std Dev", "0.3373", "MIMIC higher"],
            ["SDNN", "0.3236", "MIMIC higher"],
        ]
    )

    add_figure(
        doc,
        fig("modality_gap_ks_distances_2026.png"),
        "Figure 10. KS distances between MIMIC PERform AF and Physionet training "
        "distributions for all 8 features. All distances in the Large (>0.3) band."
    )

    add_heading(doc, "4b.2 SVM at Fixed Threshold 0.34", level=2)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Sensitivity", "100% (19/19 AF cases)"],
            ["Specificity", "12.5% (2/16 NSR correct)"],
            ["AUROC", "0.8586"],
            ["F1 Score", "0.7308"],
            ["Confusion Matrix", "TP=19, FP=14, FN=0, TN=2"],
            ["Tier 1", "PASS"],
            ["Tier 2", "FAIL (specificity 12.5%)"],
            ["Tier 3", "PASS (AUROC 0.8586 \u22650.85)"],
        ]
    )
    add_para(doc,
        "100% sensitivity with 12.5% specificity is coherent and mechanistically explained. "
        "The KS gap directly predicts this: all 8 MIMIC features are shifted toward values "
        "the model associates with Abnormal, so almost every subject crosses the "
        "ECG-calibrated threshold of 0.34 \u2014 including NSR subjects. The model does not "
        "fail to discriminate; it fails to calibrate."
    )

    add_figure(
        doc,
        fig("mimic_probability_scores_2026.png"),
        "Figure 11. SVM probability scores for MIMIC subjects at fixed threshold 0.34. "
        "AF subjects (orange) and NSR subjects (blue) are both predominantly above 0.34, "
        "with AF subjects scoring higher overall."
    )

    add_heading(doc, "4b.3 Cross-Model Comparison", level=2)
    add_para(doc,
        "All four Layer 1 models applied to MIMIC at their fixed Layer 1 thresholds. "
        "Specificity failure is systemic \u2014 confirming a modality gap problem, not a "
        "model selection artefact:"
    )
    add_table(doc,
        ["Model", "Threshold", "Sensitivity", "Specificity", "AUROC"],
        [
            ["Logistic Regression", "0.41", "100%", "18.75%", "0.7368"],
            ["Random Forest", "0.37", "100%", "6.25%", "0.8503"],
            ["XGBoost", "0.34", "100%", "12.5%", "0.8322"],
            ["SVM", "0.34", "100%", "12.5%", "0.8586"],
        ]
    )
    add_para(doc, "SVM retains the highest AUROC across all four models. Model selection confirmed correct.")

    add_heading(doc, "4b.4 Threshold Recalibration (Youden\u2019s J)", level=2)
    add_table(doc,
        ["Metric", "Fixed Threshold (0.34)", "Recalibrated (0.8424)"],
        [
            ["Sensitivity", "100%", "73.7%"],
            ["Specificity", "12.5%", "93.8%"],
            ["F1", "0.7308", "0.8235"],
        ]
    )
    add_para(doc,
        "Specificity recovers from 12.5% to 93.8% with a domain-adapted threshold. "
        "The gap between 0.34 and 0.8424 quantifies the magnitude of the modality shift. "
        "The threshold was optimised on the same 35 subjects, so stress testing is required."
    )

    add_figure(
        doc,
        fig("threshold_recalibration_2026.png"),
        "Figure 12. ROC curve for SVM on MIMIC PERform AF. Youden\u2019s J optimal "
        "threshold (0.8424) marked, showing specificity recovery from 12.5% to 93.8%."
    )

    add_heading(doc, "4b.5 Stress Testing", level=2)
    add_table(doc,
        ["Test", "Metric", "Result"],
        [
            ["Bootstrap AUROC (n=1000)", "Mean [95% CI]", "0.8631 [0.7246, 0.9720]"],
            ["Bootstrap Recalibrated Sens", "Mean [95% CI]", "80.6% [57.1%, 100.0%]"],
            ["Bootstrap Recalibrated Spec", "Mean [95% CI]", "91.3% [72.7%, 100.0%]"],
            ["LOOCV Youden\u2019s J", "Sens / Spec", "68.4% / 81.2%"],
            ["Sensitivity-Targeted LOOCV", "Sens / Spec", "78.9% / 81.2%"],
        ]
    )
    add_para(doc,
        "AUROC lower bound 0.7246 confirms discriminative signal persists even in unlucky "
        "bootstrap resamples. LOOCV specificity (81.2%) clears the pre-registered 75% "
        "criterion in both methods. Sensitivity-targeted LOOCV achieves 78.9% sensitivity "
        "\u2014 1.1 percentage points below the 80% criterion, consistent with N=19 AF "
        "leave-one-out variance."
    )

    add_figure(
        doc,
        fig("bootstrap_auroc_distribution_2026.png"),
        "Figure 13. Bootstrap AUROC distribution (n=1000 resamples) for SVM on MIMIC "
        "PERform AF. Mean 0.8631, 95% CI [0.7246, 0.9720]."
    )

    add_figure(
        doc,
        fig("loocv_comparison_2026.png"),
        "Figure 14. Comparison of Youden\u2019s J LOOCV and Sensitivity-Targeted LOOCV. "
        "Sensitivity improves 10.5pp with no specificity cost under the sensitivity-first "
        "threshold criterion."
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 4c \u2014 BeatCheck App
    # =========================================================================
    add_heading(doc, "4c. BeatCheck \u2014 Cardiac Screening Proof-of-Concept App", level=1)

    add_para(doc,
        "BeatCheck is a consumer-facing web application that applies the trained SVM model "
        "to Apple Health heart rate exports, demonstrating the research pipeline in a "
        "deployable format. It is a screening tool \u2014 not a diagnostic device \u2014 and "
        "directs users toward clinical evaluation rather than providing a diagnosis."
    )

    add_heading(doc, "4c.1 Technical Architecture", level=2)
    add_table(doc,
        ["Component", "Technology", "Notes"],
        [
            ["Backend", "FastAPI (Python)", "Serves model inference via REST API"],
            ["Frontend", "Plain HTML/CSS/JavaScript", "No framework dependencies"],
            ["ML inference", "scikit-learn SVM + StandardScaler", "Loaded from .joblib at startup"],
            ["Input formats", "Apple Health CSV, XML", "ZIP upload removed from UI"],
            ["Large file handling", "Client-side JavaScript streaming", "Full XML never transmitted"],
            ["Data window", "Most recent 90 days", "Client-side filter before upload"],
        ]
    )

    add_heading(doc, "4c.2 Processing Pipeline", level=2)
    add_para(doc,
        "Uploaded heart rate data is windowed into 30-minute overlapping windows "
        "(15-minute step, minimum 10 readings per window). Each window is scored by the "
        "frozen SVM at the domain-adapted threshold of 0.8368 (LOOCV mean threshold \u2014 "
        "chosen over the point estimate of 0.8424 for being an out-of-sample estimate). "
        "The percentage of windows flagged as Abnormal is computed and mapped to three "
        "pre-registered risk tiers:"
    )
    add_table(doc,
        ["Risk Tier", "Windows Flagged", "Interpretation", "Recommended Action"],
        [
            ["Low", "< 10%", "No consistent pattern of concern", "Continue regular monitoring"],
            ["Intermediate", "10\u201340%", "Elevated pattern detected", "Consider GP or specialist follow-up"],
            ["High", "> 40%", "Consistent Abnormal pattern", "Clinical cardiac evaluation recommended"],
        ]
    )

    add_heading(doc, "4c.3 Stress Test Results", level=2)
    add_para(doc,
        "Real Apple Health export (4.8-year dataset) processed successfully: "
        "Intermediate tier returned, 38.2% of windows flagged across 84 days analysed. "
        "602 windows processed. No crashes, no timeout errors. Client-side streaming "
        "handled the full XML export without transmitting the full file to the server."
    )

    add_figure(
        doc,
        fig("risk_tier_explanation_2026.png"),
        "Figure 15. BeatCheck risk tier framework. Three bands (Low/Intermediate/High) "
        "mapped from percentage of flagged windows, with plain-English interpretations."
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 5 \u2014 Recommendations & Next Steps
    # =========================================================================
    add_heading(doc, "5. Recommendations & Next Steps", level=1)

    add_heading(doc, "5.1 Answering the Research Question", level=2)
    add_para(doc,
        "Can machine learning models trained on clinical ECG data generalise to consumer "
        "wearable signals to detect cardiac rhythm anomalies, and are consumer wearables "
        "feasible as proxy screening tools?",
        bold=True
    )
    add_para(doc,
        "The answer is a qualified yes. The discriminative signal transfers across modalities: "
        "AUROC 0.8586 on MIMIC PERform AF confirms the model retains genuine ability to rank "
        "abnormal subjects above normal subjects using PPG-derived features. But direct "
        "threshold generalisation fails \u2014 the ECG-calibrated threshold of 0.34 produces "
        "100% sensitivity at 12.5% specificity on PPG data. This is a calibration failure, "
        "not a discrimination failure."
    )
    add_para(doc,
        "The Fahrenheit-to-Celsius analogy: a thermometer calibrated in Fahrenheit gives "
        "alarming readings when interpreted on a Celsius scale. The instrument is accurate; "
        "the interpretation framework needs adaptation. When the threshold is recalibrated "
        "for the PPG domain (0.8368 via LOOCV), sensitivity-targeted LOOCV achieves 78.9% "
        "sensitivity at 81.2% specificity \u2014 approaching the pre-registered criterion "
        "within the variance expected from N=35 leave-one-out evaluation."
    )

    add_heading(doc, "5.2 Implications for Singapore Public Health", level=2)
    add_para(doc,
        "With 49% of Singapore residents skipping annual health checkups and CVD claiming "
        "approximately 22 lives daily, the screening gap is real and consequential. "
        "Consumer wearables are worn by millions and collect continuous cardiac data without "
        "active participation. The discriminative foundation demonstrated by this project "
        "suggests the technology is viable \u2014 the engineering problem is calibration, "
        "not capability."
    )
    add_para(doc,
        "For Singapore\u2019s public health strategy, the practical path forward involves:"
    )
    add_para(doc,
        "1. PPG-native calibration: Train or fine-tune models on PPG-specific data from "
        "the target population, rather than applying ECG-trained thresholds directly. "
        "A population-representative PPG dataset with confirmed labels would dramatically "
        "improve deployability.",
        indent=True
    )
    add_para(doc,
        "2. Richer wearable data access: Partnering with wearable manufacturers to access "
        "raw beat-by-beat RR intervals (rather than summary SDNN statistics) would enable "
        "more precise HRV computation, closing part of the modality gap.",
        indent=True
    )
    add_para(doc,
        "3. Population-representative validation: MIMIC PERform AF uses critically ill ICU "
        "patients \u2014 a population that differs significantly from the general screening "
        "target. Validation on community-dwelling adults wearing consumer wearables is the "
        "necessary next step.",
        indent=True
    )
    add_para(doc,
        "4. Scope boundary: HRV-based screening can flag rhythm irregularities (AF, "
        "other arrhythmias). It cannot detect morphological abnormalities (conduction "
        "delays, ST changes) \u2014 as demonstrated by the personal ECG report. "
        "Deployment should communicate this scope explicitly to users.",
        indent=True
    )

    add_heading(doc, "5.3 Methodological Strengths", level=2)
    strengths = [
        ("Pre-registration", "Key decisions (feature set, thresholds, test direction, "
         "tier criteria) locked before analysis, preventing post-hoc rationalisation."),
        ("No data leakage", "Scaler fitted on training data only. Layer 2 data never used "
         "for training or threshold selection at the primary evaluation stage."),
        ("Transparent null result", "Apple Watch null result reported rather than "
         "suppressed. Three mechanistic explanations provided. Pivot to MIMIC documented "
         "with explicit rationale."),
        ("Discrimination vs calibration", "AUROC and sensitivity/specificity reported and "
         "interpreted separately, avoiding the common error of conflating threshold "
         "miscalibration with model failure."),
        ("Stress testing", "Bootstrap CI and LOOCV provide honest uncertainty estimates "
         "rather than point estimates that overstate precision on N=35."),
        ("Clinical ground truth", "The NHCS ECG report provides a verifiable anchor point "
         "and reveals the specific abnormality type, enabling mechanistic interpretation."),
    ]
    for title, desc in strengths:
        add_para(doc, f"{title}: {desc}", indent=True)

    add_heading(doc, "5.4 Limitations", level=2)
    limitations = [
        "Not a clinical trial and does not produce a validated medical device.",
        "Apple Watch analysis is a single-subject case study (N=1) \u2014 findings "
        "cannot be generalised.",
        "MIMIC PERform AF uses 35 critically ill ICU subjects \u2014 not representative of "
        "the general screening population.",
        "Threshold recalibration and LOOCV results computed on the same 35 subjects \u2014 "
        "indicative, not definitive.",
        "Apple Watch SE produces optical PPG, not electrical ECG. RMSSD is approximated "
        "from SDNN (\u00d70.85) due to device API limitations.",
        "The locked 8-feature HRV set detects rhythm irregularities only. Morphological "
        "abnormalities (conduction delays, ST changes) cannot be detected by any HRV-based "
        "analysis \u2014 as confirmed by the personal clinical ECG finding.",
    ]
    for lim in limitations:
        add_para(doc, f"\u2022 {lim}")

    add_heading(doc, "5.5 Acknowledgements", level=2)
    add_para(doc,
        "Dataset provided by PhysioNet and the Computing in Cardiology Challenge 2017 "
        "organisers. Personal health data collected via Apple Watch SE (1st Generation) "
        "and exported through Apple Health. MIMIC PERform AF dataset provided via Zenodo "
        "(doi.org/10.5281/zenodo.6807402). Clinical ECG report provided by the National "
        "Heart Centre Singapore (NHCS), Department of Cardiology, certified by "
        "Dr Chong Jia Yun (B.Med, FRACP), dated 17 March 2026."
    )

    # =========================================================================
    # Save
    # =========================================================================
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
